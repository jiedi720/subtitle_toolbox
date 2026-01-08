"""
文件合并模块
负责各种文档类型的合并功能，包括PDF、TXT和Word文档。
"""

import os

__all__ = [
    'run_pdf_merge_task',
    'run_txt_merge_task',
    'run_docx_merge_task',
    'run_win32_merge_task',
    'execute_merge_tasks'
]

# PDF合并所需的导入
try:
    from pypdf import PdfWriter as PdfMerger 
except ImportError:
    PdfMerger = None

# Word合并所需的导入
# 只使用python-docx进行纯Python合并
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def run_pdf_merge_task(target_dir, log_func, progress_bar, root, output_dir=None):
    """运行PDF文档合并任务
    
    合并多个PDF文档为一个。
    
    Args:
        target_dir: 目标目录
        log_func: 日志记录函数
        progress_bar: 进度条信号
        root: 根窗口
        output_dir: 输出目录
    """
    if PdfMerger is None: 
        return log_func("❌ 缺少 pypdf 库，请安装。")
    
    # 查找PDF文件
    root_files = sorted([os.path.join(target_dir, f) for f in os.listdir(target_dir) 
                        if f.lower().endswith('.pdf') and "合并" not in f])
    
    target_files = root_files if root_files else []
    save_dir = target_dir

    if not target_files:
        # 检查目标目录下的pdf子文件夹
        sub_dir = os.path.join(target_dir, "pdf")
        if os.path.exists(sub_dir):
            target_files = sorted([os.path.join(sub_dir, f) for f in os.listdir(sub_dir)
                                 if f.lower().endswith('.pdf') and "合并" not in f])
            save_dir = sub_dir

    if not target_files: 
        return log_func("❌ 未找到 PDF 文件")

    merger = PdfMerger()
    try:
        for i, f in enumerate(target_files):
                log_func(f"合并中: {os.path.basename(f)}")
                merger.append(f)
                progress_bar.emit(int((i + 1) / len(target_files) * 100))
            
        # 输出合并后的PDF
        out_path = os.path.join(save_dir, "PDF合并.pdf")
        merger.write(out_path)
        merger.close()
        log_func(f"✅ 合并成功: {out_path.replace('/', '\\')}")
    except Exception as e: 
        log_func(f"❌ 错误: {e}")
    finally: 
        progress_bar.emit(0)


def run_txt_merge_task(target_dir, log_func, progress_bar, root, output_dir=None):
    """运行TXT文档合并任务

    合并多个TXT文档为一个。

    Args:
        target_dir: 目标目录
        log_func: 日志记录函数
        progress_bar: 进度条信号
        root: 根窗口
        output_dir: 输出目录
    """
    # 查找TXT文件
    root_files = sorted([os.path.join(target_dir, f) for f in os.listdir(target_dir) 
                        if f.lower().endswith('.txt') and "合并" not in f])
    
    target_files = []
    save_dir = target_dir

    if root_files:
        target_files = root_files
    else:
        # 检查目标目录下的txt子文件夹
        sub_dir = os.path.join(target_dir, "txt")
        if os.path.exists(sub_dir):
            sub_files = sorted([os.path.join(sub_dir, f) for f in os.listdir(sub_dir)
                               if f.lower().endswith('.txt') and "合并" not in f])
            if sub_files:
                target_files = sub_files
                save_dir = sub_dir

    if not target_files:
        return log_func("❌ 未找到 TXT 文件")

    total = len(target_files)
    out_path = os.path.join(save_dir, "TXT合并.txt")

    try:
        with open(out_path, 'w', encoding='utf-8') as outfile:
            for i, fp in enumerate(target_files):
                log_func(f"合并中: {os.path.basename(fp)}")
                with open(fp, 'r', encoding='utf-8') as infile:
                    outfile.write(infile.read())
                    outfile.write("\n\n" + "="*50 + "\n\n") 
                progress_bar.emit(int((i + 1) / total * 100))
        log_func(f"✅ 合并成功: {out_path.replace('/', '\\')}")
    except Exception as e:
        log_func(f"❌ 合并失败: {e}")
    finally:
        progress_bar.emit(0)


def run_docx_merge_task(target_dir, log_func, progress_bar, root, output_dir=None):
    """使用python-docx合并Word文档
    
    通过纯Python方式合并多个Word文档为一个，不依赖Windows COM接口。
    
    Args:
        target_dir: 目标目录
        log_func: 日志记录函数
        progress_bar: 进度条对象
        root: 根窗口
        output_dir: 输出目录
    """
    if not HAS_DOCX:
        return False, "未安装 python-docx 库"
    
    root_files = sorted([os.path.join(target_dir, f) for f in os.listdir(target_dir) 
                        if f.lower().endswith('.docx') and "~$" not in f and "合并" not in f])
    
    target_files = root_files if root_files else []
    save_dir = target_dir

    if not target_files:
        # 适配新的分类层级：检测 script/word
        sub_dir = os.path.join(target_dir, "script", "word")
        if os.path.exists(sub_dir):
            target_files = sorted([os.path.join(sub_dir, f) for f in os.listdir(sub_dir)
                                 if f.lower().endswith('.docx') and "~$" not in f and "合并" not in f])
            save_dir = sub_dir
        else:
            # 检查原有的word子文件夹
            sub_dir = os.path.join(target_dir, "word")
            if os.path.exists(sub_dir):
                target_files = sorted([os.path.join(sub_dir, f) for f in os.listdir(sub_dir)
                                     if f.lower().endswith('.docx') and "~$" not in f and "合并" not in f])
                save_dir = sub_dir

    if not target_files:
        return False, "未找到 Word 文件"
    
    try:
        # 创建新文档
        merged_doc = Document()
        
        for i, fp in enumerate(target_files):
            log_func(f"合并中: {os.path.basename(fp)}")
            # 使用绝对路径并确保文件存在
            abs_path = os.path.abspath(fp)
            if os.path.exists(abs_path):
                # 打开源文档
                src_doc = Document(abs_path)
                
                # 复制所有段落
                for para in src_doc.paragraphs:
                    new_para = merged_doc.add_paragraph()
                    new_para.text = para.text
                    # 复制段落格式
                    try:
                        new_para.style = para.style
                    except AttributeError:
                        pass
                
                # 复制所有表格
                for table in src_doc.tables:
                    # 添加新表格
                    new_table = merged_doc.add_table(rows=table.rows.count, cols=table.columns.count)
                    # 复制表格内容
                    for row_idx, row in enumerate(table.rows):
                        for col_idx, cell in enumerate(row.cells):
                            new_table.cell(row_idx, col_idx).text = cell.text
                
                # 在文档之间插入分页符和新节，除了最后一个文档
                if i < len(target_files)-1:
                    merged_doc.add_page_break()
                    merged_doc.add_section()
                
                # 获取当前文档对应的节
                # 注意：第i个文件对应第i个节
                current_section = merged_doc.sections[i]
                
                # 设置节的页眉页脚不链接到前一节，确保独立
                current_section.header.is_linked_to_previous = False
                current_section.footer.is_linked_to_previous = False
                
                # 处理页眉 - 完整保留每个文档原有的页眉，并确保居中
                src_header = src_doc.sections[0].header
                merged_header = current_section.header
                
                # 清空合并文档当前节的页眉
                while len(merged_header.paragraphs) > 0:
                    merged_header._element.remove(merged_header.paragraphs[0]._element)
                
                # 复制源文档页眉的内容
                for src_para in src_header.paragraphs:
                    new_para = merged_header.add_paragraph()
                    # 设置页眉段落居中对齐
                    new_para.alignment = 1  # 1 表示居中对齐
                    for src_run in src_para.runs:
                        new_run = new_para.add_run(src_run.text)
                        # 复制基本文本格式
                        try:
                            new_run.bold = src_run.bold
                        except AttributeError:
                            pass
                        try:
                            new_run.italic = src_run.italic
                        except AttributeError:
                            pass
                        try:
                            new_run.underline = src_run.underline
                        except AttributeError:
                            pass
                
                # 处理页脚 - 完整保留每个文档原有的页脚
                src_footer = src_doc.sections[0].footer
                merged_footer = current_section.footer
                
                # 清空合并文档当前节的页脚
                while len(merged_footer.paragraphs) > 0:
                    merged_footer._element.remove(merged_footer.paragraphs[0]._element)
                
                # 复制源文档页脚的内容
                for src_para in src_footer.paragraphs:
                    new_para = merged_footer.add_paragraph()
                    for src_run in src_para.runs:
                        new_run = new_para.add_run(src_run.text)
                        # 复制基本文本格式
                        try:
                            new_run.bold = src_run.bold
                        except AttributeError:
                            pass
                        try:
                            new_run.italic = src_run.italic
                        except AttributeError:
                            pass
                        try:
                            new_run.underline = src_run.underline
                        except AttributeError:
                            pass
                
                # 尝试不同的进度条更新方式
                try:
                    # 尝试PyQt的信号方式
                    progress_bar.emit(int((i + 1) / len(target_files) * 100))
                except AttributeError:
                    try:
                        # 尝试Tkinter的set方式
                        progress_bar.set((i + 1) / len(target_files))
                        # 更新UI
                        if root:
                            root.update_idletasks()
                    except AttributeError:
                        pass
            else:
                log_func(f"❌ 文件不存在: {os.path.basename(fp)}")
        
        # 保存合并后的文档
        out_path = os.path.join(save_dir, "Word合并.docx")
        abs_out_path = os.path.abspath(out_path)
        
        # 检查输出文件是否可以被写入（检查文件是否被占用）
        if os.path.exists(abs_out_path):
            try:
                # 尝试以写入模式打开文件，检查是否被占用
                with open(abs_out_path, 'a') as f:
                    pass
            except IOError:
                log_func(f"❌ 输出文件被占用，请关闭 Word 后重试: {out_path}")
                return False, "输出文件被占用"
        
        # 保存文档
        merged_doc.save(abs_out_path)
        log_func(f"✅ 合并完成: {out_path}")
        return True, ""
    except Exception as e:
        return False, str(e)


def run_win32_merge_task(target_dir, log_func, progress_bar, root, output_dir=None):
    """合并Word文档（兼容旧接口，现在使用python-docx）
    
    直接使用python-docx合并多个Word文档为一个。
    
    Args:
        target_dir: 目标目录
        log_func: 日志记录函数
        progress_bar: 进度条对象
        root: 根窗口
        output_dir: 输出目录
    """
    # 直接调用python-docx合并方案
    success, error_msg = run_docx_merge_task(target_dir, log_func, progress_bar, root, output_dir)
    if not success:
        log_func(f"❌ 合并失败: {error_msg}")


def execute_merge_tasks(path_var, output_path_var, log_callback, update_progress, root, gui):
    """
    执行合并任务
    
    Args:
        path_var: 源目录路径
        output_path_var: 输出目录路径
        log_callback: 日志回调函数
        update_progress: 进度回调函数
        root: 根窗口对象
        gui: GUI 对象
    """
    from PySide6.QtWidgets import QMessageBox
    
    # 获取目标目录
    target_dir = path_var.strip()
    if not target_dir or not os.path.exists(target_dir):
        QMessageBox.critical(root, "错误", "请选择有效目录")
        return
    
    # 获取输出目录
    if output_path_var.strip():
        final_out = output_path_var.strip()
    else:
        final_out = target_dir
    
    # 检查Merge标签页中的复选框状态
    try:
        # 获取Merge标签页中的复选框状态
        merge_pdf = gui.MergePDF.isChecked()
        merge_word = gui.MergeWord.isChecked()
        merge_txt = gui.MergeTxt.isChecked()
        
        # 检查是否至少选中了一个合并选项
        if not merge_pdf and not merge_word and not merge_txt:
            log_callback("❌ 请至少选择一个合并选项")
            return
        
        # 根据选中的选项执行相应的合并任务
        if merge_pdf:
            run_pdf_merge_task(
                target_dir, 
                log_callback, 
                update_progress, 
                root, 
                output_dir=final_out
            )
        
        if merge_word:
            run_win32_merge_task(
                target_dir, 
                log_callback, 
                update_progress, 
                root, 
                output_dir=final_out
            )
        
        if merge_txt:
            run_txt_merge_task(
                target_dir, 
                log_callback, 
                update_progress, 
                root, 
                output_dir=final_out
            )
            
    except Exception as e:
        log_callback(f"❌ Merge模式处理失败: {e}")
